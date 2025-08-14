import os
import json
import logging
from datetime import datetime
from typing import List, Dict, Any
import time

# Google Drive API
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import io
from googleapiclient.http import MediaIoBaseDownload

# Document processing
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd

# Vector store
from pinecone import Pinecone
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Create directories if they don't exist
os.makedirs('logs', exist_ok=True)
os.makedirs('data', exist_ok=True)
os.makedirs('data/downloads', exist_ok=True)
os.makedirs('data/processed', exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/document_ingestion.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class DocumentIngestionPipeline:
    """Pipeline to ingest documents from Google Drive to Pinecone."""
    
    def __init__(self):
        self.drive_service = None
        self.pc = None
        self.embeddings = None
        self.vectorstore = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
    def setup_credentials(self):
        """Setup Google Drive credentials."""
        try:
            credentials_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE")
            if not credentials_path:
                raise Exception("GOOGLE_SERVICE_ACCOUNT_FILE not found in environment variables")
            
            if not os.path.exists(credentials_path):
                raise Exception(f"Credentials file not found at: {credentials_path}")
            
            # Load credentials
            credentials = Credentials.from_service_account_file(
                credentials_path,
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )
            
            # Build the service
            self.drive_service = build('drive', 'v3', credentials=credentials)
            logger.info("✅ Google Drive credentials setup successfully")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error setting up Google Drive credentials: {e}")
            return False
    
    def setup_pinecone(self):
        """Setup Pinecone connection."""
        try:
            pinecone_api_key = os.getenv("PINECONE_API_KEY")
            google_api_key = os.getenv("GOOGLE_API_KEY")
            index_name = os.getenv("PINECONE_INDEX_NAME", "ragn8n")
            
            if not pinecone_api_key or not google_api_key:
                raise Exception("Missing Pinecone or Google API keys")
            
            # Initialize Pinecone
            self.pc = Pinecone(api_key=pinecone_api_key)
            
            # Initialize embeddings
            self.embeddings = GoogleGenerativeAIEmbeddings(
                model="models/text-embedding-004",
                google_api_key=google_api_key
            )
            
            # Initialize vector store
            self.vectorstore = PineconeVectorStore(
                index=self.pc.Index(index_name),
                embedding=self.embeddings
            )
            
            logger.info("✅ Pinecone setup successfully")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error setting up Pinecone: {e}")
            return False
    
    def get_drive_files(self, folder_id: str) -> List[Dict]:
        """Get all files from a Google Drive folder."""
        try:
            files = []
            page_token = None
            
            while True:
                # Query for files in the folder
                query = f"'{folder_id}' in parents and trashed=false"
                results = self.drive_service.files().list(
                    q=query,
                    pageSize=100,
                    pageToken=page_token,
                    fields="nextPageToken, files(id, name, mimeType, size, modifiedTime)"
                ).execute()
                
                files.extend(results.get('files', []))
                page_token = results.get('nextPageToken')
                
                if not page_token:
                    break
            
            logger.info(f"📁 Found {len(files)} files in Google Drive folder")
            return files
            
        except HttpError as e:
            logger.error(f"❌ Error accessing Google Drive: {e}")
            return []
    
    def download_file_content(self, file_id: str, mime_type: str) -> str:
        """Download and extract text content from a file."""
        try:
            # Handle different file types
            if mime_type == 'application/pdf':
                return self._extract_pdf_content(file_id)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return self._extract_docx_content(file_id)
            elif mime_type == 'text/plain':
                return self._extract_text_content(file_id)
            elif mime_type in ['application/vnd.google-apps.document']:
                return self._extract_google_doc_content(file_id)
            elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                return self._extract_excel_content(file_id)
            else:
                logger.warning(f"⚠️ Unsupported file type: {mime_type}")
                return ""
                
        except Exception as e:
            logger.error(f"❌ Error downloading file content: {e}")
            return ""
    
    def _extract_pdf_content(self, file_id: str) -> str:
        """Extract text from PDF file."""
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_io.seek(0)
            
            # Extract text using PyPDF2
            pdf_reader = PyPDF2.PdfReader(file_io)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting PDF content: {e}")
            return ""
    
    def _extract_docx_content(self, file_id: str) -> str:
        """Extract text from DOCX file."""
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_io.seek(0)
            
            # Extract text using python-docx
            doc = DocxDocument(file_io)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting DOCX content: {e}")
            return ""
    
    def _extract_text_content(self, file_id: str) -> str:
        """Extract content from plain text file."""
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            return file_io.getvalue().decode('utf-8')
            
        except Exception as e:
            logger.error(f"❌ Error extracting text content: {e}")
            return ""
    
    def _extract_google_doc_content(self, file_id: str) -> str:
        """Extract content from Google Docs."""
        try:
            # Export Google Doc as plain text
            request = self.drive_service.files().export_media(
                fileId=file_id,
                mimeType='text/plain'
            )
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            return file_io.getvalue().decode('utf-8')
            
        except Exception as e:
            logger.error(f"❌ Error extracting Google Doc content: {e}")
            return ""
    
    def _extract_excel_content(self, file_id: str) -> str:
        """Extract content from Excel files."""
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_io.seek(0)
            
            # Read Excel file
            df = pd.read_excel(file_io, sheet_name=None)  # Read all sheets
            
            text = ""
            for sheet_name, sheet_df in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_df.to_string(index=False) + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting Excel content: {e}")
            return ""
    
    def process_document(self, file_info: Dict, content: str) -> List[Document]:
        """Process a document into chunks for vector storage."""
        if not content.strip():
            return []
        
        # Create metadata
        metadata = {
            'source': file_info['name'],
            'file_id': file_info['id'],
            'mime_type': file_info['mimeType'],
            'size': file_info.get('size', 0),
            'modified_time': file_info.get('modifiedTime', ''),
            'ingestion_time': datetime.now().isoformat()
        }
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(content)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_id'] = i
            chunk_metadata['total_chunks'] = len(chunks)
            
            documents.append(Document(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        return documents
    
    def ingest_documents(self, documents: List[Document]) -> bool:
        """Ingest documents into Pinecone."""
        try:
            if not documents:
                logger.warning("⚠️ No documents to ingest")
                return True
            
            # Add documents to vector store
            self.vectorstore.add_documents(documents)
            logger.info(f"✅ Successfully ingested {len(documents)} document chunks")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error ingesting documents: {e}")
            return False
    
    def run_ingestion(self, folder_id: str = None):
        """Run the complete document ingestion pipeline."""
        logger.info("🚀 Starting document ingestion pipeline...")
        
        # Use folder_id from parameter or environment
        if not folder_id:
            folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
        
        if not folder_id:
            logger.error("❌ No Google Drive folder ID provided")
            return False
        
        # Setup connections
        if not self.setup_credentials():
            return False
        
        if not self.setup_pinecone():
            return False
        
        # Get files from Drive
        files = self.get_drive_files(folder_id)
        if not files:
            logger.warning("⚠️ No files found in the specified folder")
            return False
        
        # Process each file
        all_documents = []
        successful_files = 0
        
        for file_info in files:
            try:
                logger.info(f"📄 Processing: {file_info['name']}")
                
                # Download and extract content
                content = self.download_file_content(file_info['id'], file_info['mimeType'])
                
                if content:
                    # Process into chunks
                    documents = self.process_document(file_info, content)
                    all_documents.extend(documents)
                    successful_files += 1
                    logger.info(f"✅ Processed {file_info['name']} -> {len(documents)} chunks")
                else:
                    logger.warning(f"⚠️ No content extracted from {file_info['name']}")
                
                # Add small delay to avoid rate limiting
                time.sleep(0.5)
                
            except Exception as e:
                logger.error(f"❌ Error processing {file_info['name']}: {e}")
                continue
        
        # Ingest all documents
        if all_documents:
            if self.ingest_documents(all_documents):
                logger.info(f"🎉 Ingestion completed successfully!")
                logger.info(f"📊 Summary: {successful_files}/{len(files)} files processed, {len(all_documents)} chunks ingested")
                return True
        
        logger.error("❌ Ingestion failed")
        return False

def main():
    """Main function to run the document ingestion."""
    print("🚀 HR Document Ingestion Pipeline")
    print("=" * 50)
    
    # Initialize pipeline
    pipeline = DocumentIngestionPipeline()
    
    # Run ingestion
    success = pipeline.run_ingestion()
    
    if success:
        print("\n✅ Document ingestion completed successfully!")
        print("Your documents are now available in the HR chatbot.")
    else:
        print("\n❌ Document ingestion failed. Please check the logs for details.")

if __name__ == "__main__":
    main()